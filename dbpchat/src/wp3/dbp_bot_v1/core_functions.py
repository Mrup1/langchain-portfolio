"""
Basic DBP Bot: Core Logic and Utilities

This file merges all core logic and helper utilities into a single flat module for maintainability.
Sections:
    - BrainAssessor: Logical vs Emotional assessment and response blending
    - Chatbot: Main orchestrator for bot sessions
    - RAGSystem: Retrieval-Augmented Generation system
    - SafetyManager: Safety filtering and event logging
    - StateTracker: Emotional state detection and tracking
    - Vector Store Utilities: Build/persist FAISS vector store

"""

# ========== Standard Library Imports ========== #
import os
import re
import json
import pickle
from enum import Enum
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime
from dotenv import load_dotenv

try:
    from langchain_groq import ChatGroq
except ImportError:
    ChatGroq = None

try:
    from langchain_google_genai import GoogleGenerativeAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document
except ImportError:
    GoogleGenerativeAIEmbeddings = None 
    FAISS = None 
    Document = None  

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# ========== Fixed Internal Imports ========== #
from .config_and_storage import Config, SafetyLevel, DatabaseManager, SafetyFilter, PromptTemplates

# ========== Section: Brain Assessor ========== #

class BrainSide(str, Enum):
    """Brain dominance sides"""
    LEFT = "left"      # Logical, analytical
    RIGHT = "right"    # Emotional, intuitive
    BALANCED = "balanced"

class BrainAssessor:
    def blend_responses(self, logical_response: str, emotional_response: str, brain_profile: dict) -> str:
        """
        Return only the dominant brain side's response based on brain_profile.
        If dominant_side is 'left', return logical; if 'right', return emotional; if 'balanced', return the more detailed/longer one.
        """
        dominant = brain_profile.get("dominant_side", "balanced")
        if dominant == "left":
            return logical_response
        elif dominant == "right":
            return emotional_response
        # If balanced, choose the longer/more detailed response
        return logical_response if len(logical_response) >= len(emotional_response) else emotional_response

    """
    Assesses user messages for brain dominance (logical vs emotional)
    and blends responses according to Dual Brain Psychology principles.
    """

    def __init__(self, config):
        """
        Initialize BrainAssessor with configuration.
        Patterns and intensity words are customizable for extensibility.
        """
        self.config = config

        # Patterns to detect logical (left brain) language
        self.logical_patterns = [
            r'\b(analyze|analysis|logic|reason|fact|data|evidence|proof|statistics)\b',
            r'\b(because|therefore|thus|hence|consequently|as a result)\b',
            r'\b(step|process|method|procedure|system|structure)\b',
            r'\b(calculate|measure|quantify|compare|evaluate)\b',
            r'\b(plan|strategy|goal|objective|target)\b',
        ]
        # Patterns to detect emotional (right brain) language
        self.emotional_patterns = [
            r'\b(feel|feeling|emotion|heart|soul|spirit)\b',
            r'\b(love|hate|fear|joy|sad|happy|angry|excited)\b',
            r'\b(intuition|instinct|gut|sense|vibe)\b',
            r'\b(creative|imagine|dream|vision|inspire)\b',
            r'\b(beautiful|amazing|wonderful|terrible|awful)\b',
        ]
        # Intensity words amplify detected patterns
        self.intensity_words = {
            'very': 1.2, 'extremely': 1.5, 'incredibly': 1.4,
            'really': 1.1, 'quite': 1.1, 'totally': 1.3,
            'absolutely': 1.4, 'completely': 1.3,
        }

    def assess_message(self, message: str) -> Dict[str, float]:
        """
        Assess a message for logical vs emotional content.
        Returns a dictionary with normalized logical_score, emotional_score, and dominant_side.
        """
        msg = message.lower()
        logical_score = sum(len(re.findall(p, msg)) for p in self.logical_patterns)
        emotional_score = sum(len(re.findall(p, msg)) for p in self.emotional_patterns)

        # Intensity multiplier
        intensity = self._get_intensity_multiplier(msg)
        logical_score *= intensity
        emotional_score *= intensity

        # Normalize scores
        total = logical_score + emotional_score
        if total > 0:
            logical_norm = logical_score / total
            emotional_norm = emotional_score / total
        else:
            logical_norm = emotional_norm = 0.5

        dominant_side = self._determine_dominance(logical_norm, emotional_norm)
        return {
            "logical_score": logical_norm,
            "emotional_score": emotional_norm,
            "dominant_side": dominant_side
        }

    @staticmethod
    def assess_and_style(message: str) -> dict:
        """
        Assess the message for brain dominance and return a dict with scores, side, and styled_text.
        """
        # Use a temporary BrainAssessor instance with default patterns
        assessor = BrainAssessor(config=None)
        scores = assessor.assess_message(message)
        dominant = scores["dominant_side"]
        # Style the message based on dominant_side
        if dominant == "left":
            styled = f"[LOGICAL] {message}"
        elif dominant == "right":
            styled = f"[EMOTIONAL] {message}"
        else:
            styled = f"[BALANCED] {message}"
        return {**scores, "styled_text": styled}

    def _get_intensity_multiplier(self, message: str) -> float:
        multiplier = 1.0
        for word, factor in self.intensity_words.items():
            if word in message:
                multiplier *= factor
        return multiplier
    
    def _determine_dominance(self, logical: float, emotional: float) -> str:
        if logical > emotional + 0.1:
            return BrainSide.LEFT
        elif emotional > logical + 0.1:
            return BrainSide.RIGHT
        else:
            return BrainSide.BALANCED

class Chatbot:
    """Main orchestrator class for Unified DBP Bot sessions."""

    def __init__(self, config: Config = Config()):
        # Load .env from the config directory
        load_dotenv(dotenv_path=Path(__file__).parent / ".env")

        self.config = config
        self.db = DatabaseManager(config.DATABASE["path"])
        self.brain_assessor = BrainAssessor(config)
        # self.state_tracker = None  
        self.safety_manager = SafetyManager(config, self.db)
        self.rag_system = RAGSystem(config)

        # Configure Groq
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables. Please check your .env file.")
        
        if ChatGroq:
            try:
                self.llm = ChatGroq(
                    model_name="gemma2-9b-it", 
                    groq_api_key=groq_api_key,
                    temperature=0.7,
                    max_tokens=1000
                )
                print(f"[INFO] Successfully initialized ChatGroq with model: gemma2-9b-it")
            except Exception as e:
                print(f"[ERROR] Failed to initialize ChatGroq: {e}")
                self.llm = None
        else:
            print("[ERROR] ChatGroq not available. Please install langchain-groq")
            self.llm = None
        self._session_id: Optional[int] = None
        self._message_sequence: int = 0
        self._conversation_history: List[Dict] = []  # raw messages

    # ------------------------------------------------------------------
    # Session management
    # ------------------------------------------------------------------
    def start_session(self, user_id: str, metadata: Optional[Dict] = None):
        """Create a new conversation session in DB and reset counters."""
        import uuid
        session_id = f"sess_{uuid.uuid4()}"
        self.db.create_session(session_id=session_id, user_id=user_id)
        self._session_id = session_id
        self._message_sequence = 0
        self._conversation_history.clear()

    # ------------------------------------------------------------------
    # Main interaction loop
    # ------------------------------------------------------------------
    def handle_user_message(self, message: str) -> str:
        # LangChain-based modular pipeline for DBP bot
        from langchain_core.runnables import RunnableLambda, RunnableMap, RunnableSequence
        import logging
        logger = logging.getLogger("DBPChatbot")
        if self._session_id is None:
            raise RuntimeError("Session not started. Call start_session() first.")
        self._message_sequence += 1
        seq = self._message_sequence
        session_id = self._session_id

        # --- Step 1: Safety Filtering ---
        def safety_step(user_msg):
            safe, filtered = self.safety_manager.check_message(
                session_id=session_id,
                message_sequence=seq,
                message=user_msg,
                is_user=True,
            )
            logger.debug(f"[SAFETY] safe={safe}, filtered={filtered}")
            return {"safe": safe, "filtered_msg": filtered if not safe else user_msg}

        # --- Step 2: State Detection ---
        def state_step(inputs):
            # Use direct state_tracker logic and log transition in DB
            
            
            # You can customize detection logic as needed
            prev_state = getattr(self, '_last_state', None)
            prev_intensity = getattr(self, '_last_intensity', 0.0)
            # Detect current state using real state detection logic
            # --- Inline state detection logic (formerly StateTracker.detect_state) ---
            message = inputs["filtered_msg"]
            context_messages = [m["user"] for m in self._conversation_history if "user" in m]
            # Example: simple keyword-based detection (expand as needed)
            stressed_patterns = [
                r"\bstress(ed|ful)?\b", r"\banxious|anxiety\b", r"\bpanic\b", r"\bworry|worried\b",
                r"\bupset\b", r"\bangry|anger\b", r"\bsad(ness)?\b", r"\bdepress(ed|ion)?\b",
                r"\boverwhelm(ed|ing)?\b", r"\bnausea|nauseous\b", r"\bkill|hurt|violence|hate\b"
            ]
            mature_patterns = [r"\bcalm\b", r"\bcollected\b", r"\brelaxed\b", r"\bsteady\b", r"\bmature\b"]
            intensity = 0.5
            detected_state = "mature"
            for pat in stressed_patterns:
                if re.search(pat, message, re.IGNORECASE):
                    detected_state = "stressed"
                    intensity = 0.8
                    break
            for pat in mature_patterns:
                if re.search(pat, message, re.IGNORECASE):
                    detected_state = "mature"
                    intensity = 0.2
                    break
            detection = {
                "to_state": detected_state,
                "to_intensity": intensity,
                "confidence": 1.0,
                "method": "keyword_pattern"
            }
            curr_state = detection.get("to_state", "mature")
            curr_intensity = detection.get("to_intensity", 0.0)
            detection_confidence = detection.get("confidence", 1.0)
            detection_method = detection.get("method", "sentiment")
            state_info = {
                "from_state": prev_state,
                "to_state": curr_state,
                "from_intensity": prev_intensity,
                "to_intensity": curr_intensity,
                "confidence": detection_confidence,
                "method": detection_method
            }
            self._last_state = curr_state
            self._last_intensity = curr_intensity

            # --- State Transition Logging (JSON) ---
            try:
                from_state = prev_state
                to_state = curr_state
                from_intensity = prev_intensity
                to_intensity = curr_intensity
                trigger_message = inputs["filtered_msg"]
                context_messages = list(self._conversation_history)
                detection_confidence = state_info.get("confidence", 1.0)
                detection_method = state_info.get("method", "default")
                debug_result = self.db.add_state_transition(
                    session_id=str(session_id),
                    from_state=from_state,
                    to_state=to_state,
                    from_intensity=from_intensity,
                    to_intensity=to_intensity,
                    trigger_message=trigger_message,
                    context_messages=context_messages,
                    detection_confidence=detection_confidence,
                    detection_method=detection_method,
                )
                import logging
                logging.getLogger("DBPChatbot").debug(f"[STATE_JSON] add_state_transition called: result={debug_result}, session_id={session_id}, from={from_state}, to={to_state}")
            except Exception as e:
                import logging
                logging.getLogger("DBPChatbot").error(f"[STATE_JSON] Error logging state transition: {e}")

            return {**inputs, "state": state_info}

        # --- Step 3: Brain Dominance ---
        def brain_step(inputs):
            # Use direct brain_assessor logic for random targeting
            
            # Randomly choose brain side for this response
            brain_result = BrainAssessor.assess_and_style(inputs["filtered_msg"])
            # Optionally, log to DB if needed (implement as desired)
            logger.debug(f"[BRAIN] {brain_result}")
            # Optionally, you can store the dominant_side for later analysis
            if hasattr(self, 'db') and hasattr(self.db, 'add_brain_dominance_record'):
                self.db.add_brain_dominance_record(
                    session_id=session_id,
                    logical_score=1.0 if brain_result["dominant_side"] == "left" else 0.0,
                    emotional_score=1.0 if brain_result["dominant_side"] == "right" else 0.0,
                    dominant_side=brain_result["dominant_side"]
                )
            # Return styled text and dominant_side for downstream steps
            return {**inputs, "brain": brain_result, "styled_msg": brain_result["styled_text"]}

        # --- Step 4: RAG Retrieval ---
        def rag_step(inputs):
            context_docs = self.rag_system.retrieve(query=inputs["filtered_msg"])
            logger.debug(f"[RAG] {context_docs}")
            return {**inputs, "context_docs": context_docs}

        # --- Step 5: Prompt Construction ---
        def prompt_step(inputs):
            # Limit context to top 2 docs and max 2000 characters
            context_docs = inputs["context_docs"][:2]
            context_str = "\n".join(context_docs)
            if len(context_str) > 2000:
                context_str = context_str[:2000] + "..."
            logical_prompt = PromptTemplates.CONVERSATION_PROMPTS["ongoing"].format(
                user_message=inputs["filtered_msg"],
                context=context_str,
                style="analytical, factual"
            )
            emotional_prompt = PromptTemplates.CONVERSATION_PROMPTS["ongoing"].format(
                user_message=inputs["filtered_msg"],
                context=context_str,
                style="empathetic, emotionally supportive"
            )
            logger.debug(f"[PROMPT] logical: {logical_prompt}")
            logger.debug(f"[PROMPT] emotional: {emotional_prompt}")
            return {**inputs, "logical_prompt": logical_prompt, "emotional_prompt": emotional_prompt}

        # --- Step 6: LLM Calls ---
        def extract_content(llm_resp):
            if isinstance(llm_resp, dict) and 'content' in llm_resp:
                return llm_resp['content']
            if hasattr(llm_resp, 'content'):
                return llm_resp.content
            if isinstance(llm_resp, str):
                return llm_resp
            return str(llm_resp)

        def llm_map(inputs):
            if self.llm is None:
                # Fallback responses when LLM is not available
                return {
                    "logical_llm": f"I understand you're saying: {inputs['filtered_msg']}. Let me think about this logically and provide a structured response.",
                    "emotional_llm": f"I hear you, and I want you to know that your feelings about '{inputs['filtered_msg']}' are completely valid. You're not alone in this.",
                    "brain": inputs["brain"],
                    "state": inputs["state"],
                    "filtered_msg": inputs["filtered_msg"]
                }
            
            try:
                logical_response = extract_content(self.llm.invoke(inputs["logical_prompt"]))
                emotional_response = extract_content(self.llm.invoke(inputs["emotional_prompt"]))
                
                return {
                    "logical_llm": logical_response,
                    "emotional_llm": emotional_response,
                    "brain": inputs["brain"],
                    "state": inputs["state"],
                    "filtered_msg": inputs["filtered_msg"]
                }
            except Exception as e:
                logger.error(f"[LLM_ERROR] {e}")
                # Fallback to simple responses
                return {
                    "logical_llm": f"I understand your message: '{inputs['filtered_msg']}'. I'm here to help you think through this step by step.",
                    "emotional_llm": f"Thank you for sharing '{inputs['filtered_msg']}' with me. I'm here to support you through whatever you're experiencing.",
                    "brain": inputs["brain"],
                    "state": inputs["state"],
                    "filtered_msg": inputs["filtered_msg"]
                }

        # --- Step 7: Blend ---
        def blend_step(inputs):
            blended = self.brain_assessor.blend_responses(
                inputs["logical_llm"], inputs["emotional_llm"], inputs["brain"]
            )
            logger.debug(f"[BLEND] {blended}")
            # Save conversation
            self._conversation_history.append({
                "user": inputs["filtered_msg"],
                "bot": blended,
                "state": inputs["state"],
                "brain": inputs["brain"]
            })
            return blended

        # --- Compose the chain ---
        dbp_chain = (
            RunnableLambda(lambda user_msg: user_msg)
            | RunnableLambda(safety_step)
            | RunnableLambda(state_step)
            | RunnableLambda(brain_step)
            | RunnableLambda(rag_step)
            | RunnableLambda(prompt_step)
            | RunnableLambda(llm_map)
            | RunnableLambda(blend_step)
        )

        return dbp_chain.invoke(message)
# =========================
# RAG System
# =========================

class RAGSystem:
    """Light-weight Retrieval-Augmented Generation helper. Currently supports FAISS."""

    def __init__(self, config: Config):
        self.config = config
        if not self.config.RAG.get("enabled", True):
            self._enabled = False
            return

        if GoogleGenerativeAIEmbeddings is None:
            raise ImportError("LangChain is required for RAG features. Install langchain>=0.1.x")

        self._enabled = True
        self.embedding_fn = GoogleGenerativeAIEmbeddings(model=config.EMBEDDING_MODEL, google_api_key=config.GEMINI_API_KEY)
        self.vector_store = None  # lazy load / build

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def build_index(self, documents: List[Dict[str, str]]):
        """Build a FAISS index from list of {id, text} dicts."""
        if not self._enabled:
            return
        docs = [Document(page_content=d["text"], metadata={"id": d["id"]}) for d in documents]
        self.vector_store = FAISS.from_documents(docs, self.embedding_fn)

    def add_documents(self, documents: List[Dict[str, str]]):
        """Add new docs to existing index or create one if absent."""
        if not self._enabled:
            return
        if self.vector_store is None:
            self.build_index(documents)
            return
        docs = [Document(page_content=d["text"], metadata={"id": d["id"]}) for d in documents]
        self.vector_store.add_documents(docs)

    def retrieve(self, query: str, top_k: int = None) -> List[str]:
        """Retrieve top-k relevant snippets for the query."""
        if not self._enabled or self.vector_store is None:
            return []
        k = top_k or self.config.RAG.get("top_k", 3)
        results = self.vector_store.similarity_search(query, k=k)
        return [res.page_content for res in results]

    def augment_prompt(self, prompt: str, context_docs: List[str]) -> str:
        """
        Concatenate retrieved context with prompt, truncating context to avoid LLM context overflow.
        Limits total prompt length to ~12000 chars (~3000 tokens). Truncates context_docs if needed.
        """
        MAX_PROMPT_CHARS = 12000
        if not context_docs:
            return prompt
        # Greedily add context_docs until limit is reached
        context_block = ""
        for doc in context_docs:
            if len(context_block) + len(doc) + len(prompt) + 32 > MAX_PROMPT_CHARS:
                break
            if context_block:
                context_block += "\n\n"
            context_block += doc
        if context_block:
            return f"Context:\n{context_block}\n\nUser:\n{prompt}"
        return prompt

# =========================
# Safety Manager
# =========================

class SafetyEventSeverity(str, Enum):
    """Enumeration for safety event severity levels"""
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"

class SafetyManager:
    """Responsible for checking messages against safety rules using SafetyFilter."""
    def __init__(self, config: Config, db_manager: DatabaseManager):
        self.config = config
        self.db = db_manager
        strictness_cfg = self.config.SAFETY["strictness"]
        self.filter = SafetyFilter(strictness=SafetyLevel(strictness_cfg.value))

    # ---------------------------------------------------------------------
    # Public API
    # ---------------------------------------------------------------------
    def check_message(
        self, session_id: int, message_sequence: int, message: str, is_user: bool = True
    ) -> Tuple[bool, str]:
        """Validate a message for safety using the configured SafetyFilter.
        Returns (is_safe, possibly_redacted_message)
        """
        if not self.config.SAFETY.get("enabled", True):
            return True, message
        filtered_msg, details = self.filter.filter_response(message)
        is_safe = not details.get("filtered", False)
        if is_safe:
            return True, message
        # Determine severity from issues (highest)
        issues: List[Dict] = details.get("issues", [])
        max_sev_int = max((issue.get("severity", 1) for issue in issues), default=1)
        severity_enum = self._map_severity(max_sev_int)
        if self.config.SAFETY.get("log_events", True):
            self._log_event(
                session_id=session_id,
                message_sequence=message_sequence,
                event_type="forbidden_term",
                severity=severity_enum,
                pattern=", ".join(issue["term"] for issue in issues),
                original_message=message,
                redacted_message=filtered_msg,
                is_user=is_user,
            )
        return False, filtered_msg

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------
    def _log_event(
        self,
        session_id: int,
        message_sequence: int,
        event_type: str,
        severity: SafetyEventSeverity,
        pattern: str,
        original_message: str,
        redacted_message: str,
        is_user: bool,
    ) -> None:
        """Persist safety event details to the DB via DatabaseManager"""
        self.db.add_safety_event(
            session_id=session_id,
            message_sequence=message_sequence,
            event_type=event_type,
            severity=severity.value,
            metadata={
                "pattern": pattern,
                "original": original_message,
                "redacted": redacted_message,
                "actor": "user" if is_user else "bot",
            },
        )

    # ------------------------------------------------------------------
    # Utility
    # ------------------------------------------------------------------
    @staticmethod
    def _map_severity(sev_int: int) -> "SafetyEventSeverity":
        if sev_int >= 4:
            return SafetyEventSeverity.HIGH
        elif sev_int >= 2:
            return SafetyEventSeverity.MEDIUM
        return SafetyEventSeverity.LOW

# ========== Section: Vector Store Utilities ========== #
"""
Utility functions to build and persist a FAISS vector store from therapy documents.
"""

# Ensure environment is loaded for API keys
load_dotenv(dotenv_path=Path(__file__).parent / ".env")

groq_api_key = os.getenv("GROQ_API_KEY")
if not groq_api_key:
    raise ValueError("GROQ_API_KEY not found in environment variables. Please check your .env file.")

try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:
    DocxDocument = None

# Helper: Load text from a .docx file
def _load_docx(path: Path) -> str:
    if DocxDocument is None:
        raise ImportError("python-docx is required for docx support.")
    doc = DocxDocument(path)
    full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(full_text)

# Helper: Load text from a .txt file
def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

# Unified document ingestion: scan conversations and Books for .txt, .docx, .pdf

def _load_pdf(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF support. Please install it.")
    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text

# Gather all .docx, .txt, and .pdf documents for ingestion from both folders
def gather_documents() -> List[Dict[str, str]]:
    """
    Scan conversations and Books directories for .docx/.txt/.pdf files and return list for indexing.
    Returns list of dicts: {"id": relative_path, "text": file_contents, "source": "conversations"|"books"}
    """
    docs = []
    sources = [
        (getattr(Config, "BOOKS_DIR", Path(__file__).parent / "Books"), "books")
    ]
    for folder, label in sources:
        folder = Path(folder)
        if not folder.exists():
            print(f"[INFO] Skipping missing folder: {folder}")
            continue
        print(f"Looking for documents in: {folder.resolve()}")
        for p in folder.glob("**/*"):
            if p.suffix.lower() in {".docx", ".txt", ".pdf"}:
                try:
                    print("Checking file:", p)
                    if p.suffix.lower() == ".docx":
                        text = _load_docx(p)
                    elif p.suffix.lower() == ".pdf":
                        text = _load_pdf(p)
                    else:
                        text = _load_txt(p)
                except Exception as exc:
                    print(f"[WARN] Skipping {p} ({exc})")
                    continue
                if text and text.strip():
                    docs.append({
                        "id": str(p.relative_to(folder)),
                        "text": text,
                        "source": label
                    })
    return docs


# Build and persist the FAISS vector store
def build_vector_store() -> None:
    """
    Build a FAISS vector store from all gathered documents and persist it to disk.
    """
    Config.ensure_dirs()
    rag = RAGSystem(Config)
    documents = gather_documents()
    if not documents:
        print("No documents found for ingestion.")
        return
    print(f"Building index for {len(documents)} documents …")
    rag.build_index(documents)
    # Persist FAISS index using FAISS's built-in save_local (avoids gRPC/pickle issues)
    index_dir = Path(Config.DATA_DIR) / "faiss_index"
    rag.vector_store.save_local(str(index_dir))
    print(f"Vector store saved to {index_dir.relative_to(Path.cwd())}")

if __name__ == "__main__":
    build_vector_store()